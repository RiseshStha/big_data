"""Post-correction verification of document images and figure captions."""
import docx
import os
import re
import hashlib

DOC_PATH = r"D:/Data Science/Big Data and Data Visualization/Assignment/Documents/Big Data only Intro to ref.docx"
PLOT_DIR = r"D:/Data Science/Big Data and Data Visualization/Assignment/Project/Outputs/Plots"
TABLEAU_DIR = r"D:/Data Science/Big Data and Data Visualization/Assignment/Project/extracted_images"

EXPECTED = {
    "rId7": (PLOT_DIR, "model_1_confusion_matrix.png"),
    "rId8": (PLOT_DIR, "model_2_class_accuracy.png"),
    "rId9": (PLOT_DIR, "model_3_feature_importance.png"),
    "rId10": (PLOT_DIR, "model_4_regression_predictions.png"),
    "rId11": (PLOT_DIR, "model_5_lda_elbow.png"),
    "rId12": (PLOT_DIR, "model_6_lda_topics_heatmap.png"),
    "rId13": (PLOT_DIR, "model_7_topic_dist.png"),
    "rId14": (TABLEAU_DIR, "image10.png"),
    "rId15": (TABLEAU_DIR, "image11.png"),
    "rId16": (TABLEAU_DIR, "image12.png"),
    "rId17": (TABLEAU_DIR, "image13.png"),
    "rId18": (TABLEAU_DIR, "image14.png"),
    "rId19": (TABLEAU_DIR, "image15.png"),
}


def md5(data):
    return hashlib.md5(data).hexdigest()


def main():
    doc = docx.Document(DOC_PATH)
    print("=== IMAGE HASH VERIFICATION ===")
    all_ok = True
    for rid, (img_dir, fname) in EXPECTED.items():
        expected_path = os.path.join(img_dir, fname)
        with open(expected_path, "rb") as f:
            expected_hash = md5(f.read())
        if rid in doc.part.related_parts:
            actual_hash = md5(doc.part.related_parts[rid]._blob)
            match = expected_hash == actual_hash
            status = "OK" if match else "MISMATCH"
            if not match:
                all_ok = False
            print(f"  [{status}] {rid} == {fname}")
        else:
            all_ok = False
            print(f"  [MISSING] {rid}")

    print()
    print("=== FIGURE CAPTIONS (in order) ===")
    fig_num = 0
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Figure 2."):
            fig_num += 1
            print(f"  P#{i:03d}: {p.text.strip()[:90]}")

    print()
    print("=== TABLE OF FIGURES (P#13-P#25) ===")
    for i in range(13, 26):
        if i < len(doc.paragraphs):
            print(f"  P#{i:03d}: {doc.paragraphs[i].text.strip()[:90]}")

    print()
    print("=== TOC HEADING ===")
    print(f"  P#009: {doc.paragraphs[9].text.strip()}")
    print(f"  P#012: {doc.paragraphs[12].text.strip()}")

    print()
    print("=== REFERENCES (first 3) ===")
    for i in range(124, 127):
        print(f"  P#{i}: {doc.paragraphs[i].text.strip()[:80]}...")

    print()
    print("=== OVERALL ===")
    print("  ALL IMAGES MATCH:" if all_ok else "  SOME IMAGES DO NOT MATCH")


if __name__ == "__main__":
    main()
