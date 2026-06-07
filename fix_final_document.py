"""
fix_final_document.py

Comprehensive correction script for the Big Data academic report.
Performs:
1. Image replacement (sections 2.3-2.5 with new Python plots)
2. Figure caption renumbering (sequential order of appearance)
3. British English spelling consistency
4. Specific text corrections
5. Harvard reference formatting (alphabetically sorted)
6. Table of Contents field insertion
7. Table of Figures insertion
8. Heading style verification
9. Page break before main body

IMPORTANT: Close Microsoft Word before running this script!
"""

import docx
import shutil
import os
import sys
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# PATHS
# ============================================================
DOC_DIR = "D:/Data Science/Big Data and Data Visualization/Assignment/Documents"
DOC_PATH = os.path.join(DOC_DIR, "Big Data only Intro to ref.docx")
PLOT_DIR = "D:/Data Science/Big Data and Data Visualization/Assignment/Project/Outputs/Plots"
TABLEAU_DIR = "D:/Data Science/Big Data and Data Visualization/Assignment/Project/extracted_images"
BACKUP_PATH = os.path.join(DOC_DIR, "Big Data only Intro to ref_pre_final.docx")


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def insert_toc_field(paragraph):
    """Insert a Word TOC field code into a paragraph."""
    p = paragraph._p

    r1 = OxmlElement('w:r')
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    r1.append(fld1)
    p.append(r1)

    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instr)
    p.append(r2)

    r3 = OxmlElement('w:r')
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'separate')
    r3.append(fld2)
    p.append(r3)

    r4 = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i_elem = OxmlElement('w:i')
    rPr.append(i_elem)
    r4.append(rPr)
    t = OxmlElement('w:t')
    t.text = "[Right-click here and select Update Field to generate Table of Contents]"
    r4.append(t)
    p.append(r4)

    r5 = OxmlElement('w:r')
    fld3 = OxmlElement('w:fldChar')
    fld3.set(qn('w:fldCharType'), 'end')
    r5.append(fld3)
    p.append(r5)


def add_page_break_to_paragraph(paragraph):
    """Add a page break at the end of a paragraph."""
    run = paragraph.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def clear_paragraph_text(paragraph):
    """Clear all text from a paragraph while preserving XML structure."""
    for run in paragraph.runs:
        run.text = ""


def set_paragraph_heading(paragraph, text, font_size=16, bold=True, center=True):
    """Set paragraph as a heading with specific formatting."""
    clear_paragraph_text(paragraph)
    if len(paragraph.runs) > 0:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run()
    run.text = text
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_paragraph_entry(paragraph, text, font_size=12, font_name="Times New Roman"):
    """Set paragraph as a list entry."""
    clear_paragraph_text(paragraph)
    if len(paragraph.runs) > 0:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run()
    run.text = text
    run.bold = False
    run.font.size = Pt(font_size)
    run.font.name = font_name


def contains_url(text):
    """Check if text contains a URL."""
    return "http" in text or "www." in text or ".org" in text or ".com" in text


# ============================================================
# MAIN SCRIPT
# ============================================================

def main():
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

    print("=" * 70)
    print("  COMPREHENSIVE DOCUMENT CORRECTION SCRIPT")
    print("  Big Data only Intro to ref.docx")
    print("=" * 70)

    # ----------------------------------------------------------
    # STEP 0: BACKUP
    # ----------------------------------------------------------
    if not os.path.exists(DOC_PATH):
        print(f"ERROR: Document not found at: {DOC_PATH}")
        return

    print(f"\n[BACKUP] Creating backup: {BACKUP_PATH}")
    shutil.copy2(DOC_PATH, BACKUP_PATH)

    # ----------------------------------------------------------
    # STEP 1: LOAD DOCUMENT
    # ----------------------------------------------------------
    print("[LOAD] Loading document...")
    doc = docx.Document(DOC_PATH)
    total_para = len(doc.paragraphs)
    print(f"  Total paragraphs: {total_para}")

    # ----------------------------------------------------------
    # PART 1: UPDATE IMAGES
    # Sections 2.3-2.5: Python-generated plots
    # Section 2.6: Original Tableau screenshots from backup
    # ----------------------------------------------------------
    print("\n" + "-" * 50)
    print("PART 1: UPDATING IMAGES")
    print("-" * 50)

    image_updates = {
        # Section 2.3 (Random Forest classification)
        "rId7": (PLOT_DIR, "model_1_confusion_matrix.png"),       # Figure 2.3 -> 2.1
        "rId8": (PLOT_DIR, "model_2_class_accuracy.png"),          # Figure 2.2
        "rId9": (PLOT_DIR, "model_3_feature_importance.png"),      # Figure 2.1 -> 2.3
        # Section 2.4 (Linear Regression)
        "rId10": (PLOT_DIR, "model_4_regression_predictions.png"),  # Figure 2.4
        # Section 2.5 (LDA topic modelling)
        "rId11": (PLOT_DIR, "model_5_lda_elbow.png"),              # Figure 2.7 -> 2.5
        "rId12": (PLOT_DIR, "model_6_lda_topics_heatmap.png"),     # Figure 2.6
        "rId13": (PLOT_DIR, "model_7_topic_dist.png"),             # Figure 2.5 -> 2.7
        # Section 2.6 (Tableau visualisation — backup screenshots)
        "rId14": (TABLEAU_DIR, "image10.png"),                     # Growth over time
        "rId15": (TABLEAU_DIR, "image11.png"),                     # Top 20 tags by avg score
        "rId16": (TABLEAU_DIR, "image12.png"),                     # Top 20 tags distribution
        "rId17": (TABLEAU_DIR, "image13.png"),                     # Score category distribution
        "rId18": (TABLEAU_DIR, "image14.png"),                     # Closed vs open
        "rId19": (TABLEAU_DIR, "image15.png"),                     # Yearly growth trends
    }
    for rid, (img_dir, fname) in image_updates.items():
        fpath = os.path.join(img_dir, fname)
        if rid in doc.part.related_parts and os.path.exists(fpath):
            with open(fpath, "rb") as f:
                doc.part.related_parts[rid]._blob = f.read()
            print(f"  [OK] {rid} -> {fname}")
        else:
            print(f"  [WARN] Could not update {rid} (rid_exists={rid in doc.part.related_parts}, file_exists={os.path.exists(fpath)})")

    # ----------------------------------------------------------
    # PART 2: RENUMBER FIGURE CAPTIONS (sequential order)
    # ----------------------------------------------------------
    print("\n" + "-" * 50)
    print("PART 2: RENUMBERING FIGURES (sequential order)")
    print("-" * 50)

    # Map: paragraph_index -> (old_prefix, new_prefix)
    # Ordered by paragraph index to avoid collision
    figure_renames = [
        (52, "Figure 2.3:", "Figure 2.1:"),
        # 54: "Figure 2.2:" stays same
        (56, "Figure 2.1:", "Figure 2.3:"),
        # 63: "Figure 2.4:" stays same
        (70, "Figure 2.7:", "Figure 2.5:"),
        # 72: "Figure 2.6:" stays same
        (74, "Figure 2.5:", "Figure 2.7:"),
        (81, "Figure 2.11:", "Figure 2.8:"),
        (82, "Figure 2.10:", "Figure 2.9:"),
        (84, "Figure 2.9:", "Figure 2.10:"),
        (86, "Figure 2.8:", "Figure 2.11:"),
        (89, "Figure 2.13:", "Figure 2.12:"),
        (91, "Figure 2.12:", "Figure 2.13:"),
    ]

    for idx, old_prefix, new_prefix in figure_renames:
        if idx < total_para:
            p = doc.paragraphs[idx]
            if old_prefix in p.text:
                replaced = False
                for run in p.runs:
                    if old_prefix in run.text:
                        run.text = run.text.replace(old_prefix, new_prefix)
                        print(f"  [OK] P#{idx}: {old_prefix} -> {new_prefix}")
                        replaced = True
                        break
                if not replaced:
                    print(f"  [WARN] P#{idx}: Found '{old_prefix}' in paragraph text but not in any run")
            else:
                print(f"  [WARN] P#{idx}: Expected '{old_prefix}' but found: '{p.text[:50]}'")

    # Also fix capitalisation: "elbow curves" -> "Elbow Curves" in Figure 2.5 caption
    if 70 < total_para:
        p = doc.paragraphs[70]
        for run in p.runs:
            if "elbow curves" in run.text:
                run.text = run.text.replace("elbow curves", "Elbow Curves")
                print(f"  [OK] P#70: Fixed 'elbow curves' -> 'Elbow Curves'")

    # ----------------------------------------------------------
    # PART 3: BRITISH ENGLISH SPELLING
    # ----------------------------------------------------------
    print("\n" + "-" * 50)
    print("PART 3: FIXING BRITISH ENGLISH SPELLING")
    print("-" * 50)

    spelling_fixes = [
        ("utilized", "utilised"),
        ("utilizes", "utilises"),
        ("summarizes", "summarises"),
        ("optimized", "optimised"),
        ("penalize", "penalise"),
        ("penalized", "penalised"),
        ("synthesize", "synthesise"),
        ("synthesized", "synthesised"),
        ("Visualization", "Visualisation"),
        ("visualization", "visualisation"),
        ("inner joined", "inner-joined"),
    ]

    fix_count = 0
    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            if contains_url(run.text):
                continue
            for old, new in spelling_fixes:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    fix_count += 1
    print(f"  Applied {fix_count} spelling correction(s)")

    # ----------------------------------------------------------
    # PART 4: SPECIFIC TEXT CORRECTIONS
    # ----------------------------------------------------------
    print("\n" + "-" * 50)
    print("PART 4: SPECIFIC TEXT CORRECTIONS")
    print("-" * 50)

    # Fix P#029: Remove informal "(5 v's)" if present
    if 29 < total_para:
        p = doc.paragraphs[29]
        for run in p.runs:
            if "(5 v" in run.text.lower():
                run.text = re.sub(r"\s*\(5\s+v['\u2019]s?\)", "", run.text)
                print("  [OK] P#029: Removed informal '(5 v's)' from Introduction")
                break

    # Fix P#043: Add missing article "a" before "master-worker"
    if 43 < total_para:
        p = doc.paragraphs[43]
        for run in p.runs:
            if "operates on master-worker" in run.text:
                run.text = run.text.replace("operates on master-worker", "operates on a master-worker")
                print("  [OK] P#043: Added missing article 'a' before 'master-worker'")
                break

    # ----------------------------------------------------------
    # PART 5: FIX REFERENCES (Harvard format, alphabetical)
    # ----------------------------------------------------------
    print("\n" + "-" * 50)
    print("PART 5: FIXING REFERENCES (Harvard format)")
    print("-" * 50)

    sorted_refs = [
        "Apache Spark (2024) Machine learning library (MLlib) guide. Available at: https://spark.apache.org/docs/latest/ml-guide.html (Accessed: 28 May 2026).",
        "Blei, D.M., Ng, A.Y. and Jordan, M.I. (2003) \u2018Latent Dirichlet Allocation\u2019, Journal of Machine Learning Research, 3, pp. 993\u20131022.",
        "Breiman, L. (2001) \u2018Random Forests\u2019, Machine Learning, 45(1), pp. 5\u201332.",
        "Build and Debug (2024) Creating a Spark session in PySpark: A step-by-step guide with real-time scenarios. Medium. Available at: https://medium.com/@BuildandDebug/creating-a-spark-session-in-pyspark-a-step-by-step-guide-with-real-time-scenarios-55a64dac2a79 (Accessed: 28 May 2026).",
        "Bulut, O. and Desjardins, C. (2021) \u2018Visualising big data\u2019, in Exploring, Visualising, and Modelling Big Data with R. Available at: https://okanbulut.github.io/bigdata/visualizing-big-data.html (Accessed: 28 May 2026).",
        "DataCamp (2024) PySpark tutorial: Getting started with PySpark. Available at: https://www.datacamp.com/tutorial/pyspark-tutorial-getting-started-with-pyspark (Accessed: 28 May 2026).",
        "IBM (2023) What are the 5 Vs of big data? Available at: https://www.ibm.com/think/topics/big-data (Accessed: 28 May 2026).",
        "Kaggle (2019) StackSample: 10% of Stack Overflow Q&A. Available at: https://www.kaggle.com/datasets/stackoverflow/stacksample (Accessed: 28 May 2026).",
        "Kavlakoglu, E. (2025) What is latent Dirichlet allocation? IBM. Available at: https://www.ibm.com/think/topics/latent-dirichlet-allocation (Accessed: 28 May 2026).",
        "Stack Overflow (2026) Stack Exchange Data Dump. Available at: https://archive.org/details/stackexchange (Accessed: 28 May 2026).",
        "Tableau (2024) Data visualisation tips and best practices. Available at: https://www.tableau.com/visualization/data-visualization-best-practices (Accessed: 28 May 2026).",
        "Use Data To Lead (2023) Five principles of data visualisation. Available at: https://www.usedatatolead.com/en/post/funf-prinzipien-der-datenvisualisierung (Accessed: 28 May 2026).",
        "Williams, D. (2026) Five steps to tackle big graph data visualisation. Cambridge Intelligence. Available at: https://cambridge-intelligence.com/blog/big-graph-data-visualization/ (Accessed: 28 May 2026).",
    ]

    ref_start_idx = 124
    for i, ref_text in enumerate(sorted_refs):
        p_idx = ref_start_idx + i
        if p_idx < total_para:
            p = doc.paragraphs[p_idx]
            # Preserve first run formatting, clear the rest
            if len(p.runs) > 0:
                p.runs[0].text = ref_text
                for j in range(1, len(p.runs)):
                    p.runs[j].text = ""
            else:
                run = p.add_run(ref_text)
                run.font.size = Pt(12)
                run.font.name = "Times New Roman"
            print(f"  [OK] Ref {i+1}/{len(sorted_refs)}: {ref_text[:55]}...")
        else:
            print(f"  [WARN] P#{p_idx} does not exist! Cannot write reference.")

    # Clear remaining old reference paragraphs (P#137 to P#140)
    ref_end_idx = ref_start_idx + len(sorted_refs)
    for p_idx in range(ref_end_idx, 141):
        if p_idx < total_para:
            p = doc.paragraphs[p_idx]
            if p.text.strip():
                clear_paragraph_text(p)
                print(f"  [OK] P#{p_idx}: Cleared old reference text")

    # ----------------------------------------------------------
    # PART 6: TABLE OF CONTENTS & TABLE OF FIGURES
    # ----------------------------------------------------------
    print("\n" + "-" * 50)
    print("PART 6: ADDING TABLE OF CONTENTS & TABLE OF FIGURES")
    print("-" * 50)

    # Layout for blank paragraphs P#009-P#027:
    # P#009: [has page break] + "Table of Contents" heading
    # P#010: TOC field
    # P#011: blank (TOC expansion)
    # P#012: "Table of Figures" heading
    # P#013-P#025: 13 figure entries
    # P#026: blank
    # P#027: [page break before Introduction]

    figure_list = [
        "Figure 2.1: Random Forest Classification Confusion Matrix Heatmap",
        "Figure 2.2: Per-Class Accuracy by Tag Category",
        "Figure 2.3: Random Forest Feature Importances for Top NLP Terms",
        "Figure 2.4: Linear Regression Log Score \u2013 Actual vs. Predicted Scatter Plot",
        "Figure 2.5: LDA Model Selection Elbow Curves for Log Perplexity and Log Likelihood",
        "Figure 2.6: Latent Dirichlet Allocation Topic-Word Weight Heatmap",
        "Figure 2.7: Document Counts Distribution across Discovered LDA Topics",
        "Figure 2.8: Stack Overflow Platform Question Growth Over Time (2008\u20132016)",
        "Figure 2.9: Top 20 Programming Tags Sorted by Average Upvote Score",
        "Figure 2.10: Distribution of Top 20 Programming Language Tags by Question Volume",
        "Figure 2.11: Stack Overflow Questions Score Category Distribution",
        "Figure 2.12: Closed vs. Open Questions Volume and Average Score Comparison",
        "Figure 2.13: Yearly Growth Trends of Top 5 Programming Tags",
    ]

    # P#009: Table of Contents heading
    if 9 < total_para:
        p9 = doc.paragraphs[9]
        run = p9.add_run("Table of Contents")
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "Times New Roman"
        p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print("  [OK] P#009: Added 'Table of Contents' heading")

    # P#010: TOC field
    if 10 < total_para:
        p10 = doc.paragraphs[10]
        insert_toc_field(p10)
        print("  [OK] P#010: Inserted TOC field code")

    # P#011: blank (leave as TOC expansion buffer)

    # P#012: Table of Figures heading
    if 12 < total_para:
        set_paragraph_heading(doc.paragraphs[12], "Table of Figures", font_size=18)
        print("  [OK] P#012: Added 'Table of Figures' heading")

    # P#013-P#025: Figure entries
    for i, fig_text in enumerate(figure_list):
        p_idx = 13 + i
        if p_idx < total_para:
            set_paragraph_entry(doc.paragraphs[p_idx], fig_text, font_size=11)
            print(f"  [OK] P#{p_idx}: {fig_text[:55]}...")

    # P#027: Add page break before "1. Introduction"
    if 27 < total_para:
        p27 = doc.paragraphs[27]
        clear_paragraph_text(p27)
        add_page_break_to_paragraph(p27)
        print("  [OK] P#027: Added page break before Introduction")

    # ----------------------------------------------------------
    # PART 7: VERIFY AND FIX HEADING STYLES
    # ----------------------------------------------------------
    print("\n" + "-" * 50)
    print("PART 7: VERIFYING HEADING STYLES (for TOC)")
    print("-" * 50)

    heading_checks = [
        (28, "Heading 1", "1. Introduction"),
        (36, "Heading 1", "2. Implementation"),
        (37, "Heading 2", "2.1 Dataset Description"),
        (41, "Heading 2", "2.2 Environment"),
        (44, "Heading 2", "2.3 Task 1"),
        (58, "Heading 2", "2.4 Task 2"),
        (64, "Heading 2", "2.5 Task 3"),
        (77, "Heading 2", "2.6 Tableau"),
        (92, "Heading 1", "3. Discussion"),
        (98, "Heading 1", "4. Conclusion"),
        (123, "Heading 1", "5. References"),
        (141, "Heading 1", "6. Appendix"),
    ]

    for p_idx, expected_style, label in heading_checks:
        if p_idx < total_para:
            p = doc.paragraphs[p_idx]
            current_style = p.style.name if p.style else "None"
            if current_style != expected_style:
                try:
                    p.style = doc.styles[expected_style]
                    print(f"  [FIXED] P#{p_idx} '{label}': {current_style} -> {expected_style}")
                except KeyError:
                    print(f"  [WARN] Style '{expected_style}' not found. P#{p_idx} has '{current_style}'")
            else:
                print(f"  [OK] P#{p_idx} '{label}': {current_style}")

    # ----------------------------------------------------------
    # SAVE
    # ----------------------------------------------------------
    print("\n" + "=" * 70)
    try:
        doc.save(DOC_PATH)
        print("  DOCUMENT SAVED SUCCESSFULLY!")
        print(f"  Path: {DOC_PATH}")
        print(f"  Backup: {BACKUP_PATH}")
    except PermissionError:
        alt_path = os.path.join(DOC_DIR, "Big Data only Intro to ref_FIXED.docx")
        try:
            doc.save(alt_path)
            print(f"  Original file is LOCKED by Word.")
            print(f"  SAVED TO ALTERNATE PATH: {alt_path}")
            print(f"  Please close Word and rename this file.")
        except Exception as e:
            print(f"  ERROR: Could not save: {e}")
            print("  Please CLOSE Microsoft Word and run this script again.")
    except Exception as e:
        print(f"  ERROR: {e}")

    print("=" * 70)
    print("\n  NEXT STEPS FOR YOU:")
    print("  1. Open the document in Microsoft Word")
    print("  2. Press Ctrl+A to select all content")
    print("  3. Press F9 to update all fields (Table of Contents)")
    print("  4. Click 'Yes' to update the Table of Contents")
    print("  5. Review the document and save (Ctrl+S)")
    print("=" * 70)


if __name__ == "__main__":
    main()
