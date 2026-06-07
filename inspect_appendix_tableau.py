"""Extract and inspect Tableau images from appendix vs section 2.6."""
import docx
import os
import re
import hashlib

DOC = r"D:/Data Science/Big Data and Data Visualization/Assignment/Documents/Big Data only Intro to ref.docx"
OUT = r"D:/Data Science/Big Data and Data Visualization/Assignment/Project/appendix_tableau_extracted"
PLOT_DIR = r"D:/Data Science/Big Data and Data Visualization/Assignment/Project/Outputs/Plots"
BACKUP_DIR = r"D:/Data Science/Big Data and Data Visualization/Assignment/Project/extracted_images"

os.makedirs(OUT, exist_ok=True)


def md5(data):
    return hashlib.md5(data).hexdigest()


def main():
    doc = docx.Document(DOC)

    print("=== SECTION 2.6 FIGURES ===")
    for i in range(77, 92):
        p = doc.paragraphs[i]
        rids = re.findall(r'r:embed="(rId\d+)"', p._element.xml)
        t = p.text.strip()
        if rids or t.startswith("Figure 2."):
            print(f"P#{i:03d} rIds={rids} | {t[:90]}")

    print("\n=== APPENDIX TABLEAU SECTION ===")
    appendix_rids = []
    for i in range(141, 200):
        if i >= len(doc.paragraphs):
            break
        p = doc.paragraphs[i]
        rids = re.findall(r'r:embed="(rId\d+)"', p._element.xml)
        t = p.text.strip()
        if rids or "Tableau" in t or "Dashboard" in t or t.startswith("Figure"):
            print(f"P#{i:03d} rIds={rids} | {t[:90]}")
            appendix_rids.extend(rids)

    print("\n=== EXTRACT APPENDIX IMAGES (first 20 unique rIds in appendix area) ===")
    seen = set()
    extracted = []
    for i in range(141, 200):
        if i >= len(doc.paragraphs):
            break
        p = doc.paragraphs[i]
        rids = re.findall(r'r:embed="(rId\d+)"', p._element.xml)
        for rid in rids:
            if rid in seen:
                continue
            seen.add(rid)
            if rid in doc.part.related_parts:
                part = doc.part.related_parts[rid]
                if "image" in getattr(part, "content_type", ""):
                    fname = f"appendix_{rid}.png"
                    path = os.path.join(OUT, fname)
                    with open(path, "wb") as f:
                        f.write(part._blob)
                    h = md5(part._blob)
                    extracted.append((rid, fname, h, i))
                    print(f"  {rid} from P#{i} -> {fname} ({len(part._blob)} bytes) hash={h[:12]}")

    print("\n=== COMPARE SECTION 2.6 rIds WITH APPENDIX / PLOTS / BACKUP ===")
    section26 = {
        "rId14": "Figure 2.8 Growth",
        "rId15": "Figure 2.9 Avg Score",
        "rId16": "Figure 2.10 Top Tags",
        "rId17": "Figure 2.11 Score Dist",
        "rId18": "Figure 2.12 Closed vs Open",
        "rId19": "Figure 2.13 Trends",
    }
    eda_map = {
        "rId14": "eda_5_growth.png",
        "rId15": "eda_3_avg_score.png",
        "rId16": "eda_2_top_tags.png",
        "rId17": "eda_1_score_dist.png",
        "rId18": "eda_7_closed_open.png",
        "rId19": "eda_6_trends.png",
    }
    backup_map = {
        "rId14": "image10.png",
        "rId15": "image11.png",
        "rId16": "image12.png",
        "rId17": "image13.png",
        "rId18": "image14.png",
        "rId19": "image15.png",
    }

    for rid, label in section26.items():
        if rid not in doc.part.related_parts:
            print(f"{rid} ({label}): NOT IN DOC")
            continue
        blob = doc.part.related_parts[rid]._blob
        h = md5(blob)

        matches = []
        for eda in eda_map.values():
            p = os.path.join(PLOT_DIR, eda)
            if os.path.exists(p):
                with open(p, "rb") as f:
                    if md5(f.read()) == h:
                        matches.append(f"generate_plots:{eda}")

        for img in backup_map.values():
            p = os.path.join(BACKUP_DIR, img)
            if os.path.exists(p):
                with open(p, "rb") as f:
                    if md5(f.read()) == h:
                        matches.append(f"backup:{img}")

        for arid, fname, ah, pi in extracted:
            if ah == h:
                matches.append(f"appendix:{arid}@P#{pi}")

        print(f"{rid} ({label}): matches {matches or ['UNKNOWN SOURCE']}")


if __name__ == "__main__":
    main()
