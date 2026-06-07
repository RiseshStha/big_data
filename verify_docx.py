"""Verify current document structure, rIds, and figure captions."""
import docx
import os
import re

DOC_PATH = r"D:/Data Science/Big Data and Data Visualization/Assignment/Documents/Big Data only Intro to ref.docx"
PLOT_DIR = r"D:/Data Science/Big Data and Data Visualization/Assignment/Project/Outputs/Plots"
EXTRACTED_DIR = r"D:/Data Science/Big Data and Data Visualization/Assignment/Project/extracted_images"


def main():
    doc = docx.Document(DOC_PATH)
    print("Total paragraphs:", len(doc.paragraphs))
    print("Total tables:", len(doc.tables))
    print()

    print("=== IMAGE RELATIONSHIPS ===")
    for rid in sorted(
        doc.part.related_parts.keys(),
        key=lambda x: int(re.search(r"\d+", x).group()) if re.search(r"\d+", x) else 0,
    ):
        if rid.startswith("rId"):
            num = int(rid[3:])
            if num <= 25:
                part = doc.part.related_parts[rid]
                ctype = getattr(part, "content_type", "unknown")
                if "image" in ctype:
                    print(f"  {rid}: {ctype}")

    print()
    print("=== PARAGRAPHS WITH IMAGES OR FIGURE CAPTIONS ===")
    for i, p in enumerate(doc.paragraphs):
        xml = p._element.xml
        has_img = "drawing" in xml or "pict" in xml or "blip" in xml
        text = p.text.strip()
        if has_img or text.startswith("Figure 2."):
            style = p.style.name if p.style else "None"
            img_rids = re.findall(r'r:embed="(rId\d+)"', xml)
            print(f"P#{i:03d} [{style}] rIds={img_rids} | {text[:100]}")

    print()
    print("=== KEY HEADINGS ===")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if re.match(r"^[1-6]\.", t) or re.match(r"^2\.\d", t):
            style = p.style.name if p.style else "None"
            print(f"P#{i:03d} [{style}] {t[:80]}")

    print()
    print("=== PLOT FILES AVAILABLE ===")
    for f in sorted(os.listdir(PLOT_DIR)):
        if f.endswith(".png"):
            print(f"  {f}")

    print()
    print("=== TABLEAU BACKUP IMAGES (extracted_images) ===")
    for f in sorted(os.listdir(EXTRACTED_DIR)):
        if f.startswith("image") and f.endswith(".png"):
            print(f"  {f}")

    print()
    print("=== FILE LOCK TEST ===")
    try:
        with open(DOC_PATH, "r+b"):
            print("  Document is NOT locked (writable)")
    except PermissionError:
        print("  Document IS LOCKED by Word or another process")


if __name__ == "__main__":
    main()
